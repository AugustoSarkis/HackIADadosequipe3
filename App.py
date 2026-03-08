import streamlit as st
import os
from google import genai
from PyPDF2 import PdfReader
from docx import Document
import io
import time

# --- 1. Configuração da Página ---
st.set_page_config(
    page_title="Assistente de Diagnóstico Comercial", 
    page_icon="💼", 
    layout="wide",
    initial_sidebar_state="expanded"
)

PRECOS_MODELOS = {
    "gemini-2.5-flash-lite": {"in": 0.075, "out": 0.30},
    "gemini-2.5-flash": {"in": 0.075, "out": 0.30},
    "gemini-2.5-pro": {"in": 1.25, "out": 5.00},
}

# --- 2. Inicialização do Cliente API ---
@st.cache_resource
def inicializar_cliente():
    try:
        os.environ["GEMINI_API_KEY"] = st.secrets["GEMINI_API_KEY"]
        return genai.Client()
    except KeyError:
        return None

client = inicializar_cliente()

# --- 3. Controle de Estado (Session State) ---
if "resultados_processados" not in st.session_state:
    st.session_state.resultados_processados = {}
if "mostrar_animacao" not in st.session_state:
    st.session_state.mostrar_animacao = False

def limpar_memoria():
    st.session_state.resultados_processados = {}
    st.toast("Memória limpa com sucesso!", icon="🧹") # Micro-interação elegante

# --- 4. Funções Modulares ---
def extrair_texto_pdf(arquivo):
    leitor_pdf = PdfReader(arquivo)
    texto = "".join([pagina.extract_text() or "" for pagina in leitor_pdf.pages])
    return texto.strip()

def gerar_prompt_multiplas_atas(texto_transcricao):
    # Prompt otimizado para forçar a IA a usar HTML/Emojis no Score
    return f"""
    Imagine que você é um revisor técnico.

    Vou te enviar anotações brutas que PODEM conter transcrições de UMA OU MÚLTIPLAS reuniões comerciais diferentes.

    Sua tarefa tem duas etapas:
    1. Identifique e separe cada reunião distinta presente no texto.
    2. Para CADA reunião identificada, interprete as anotações e transforme-as em uma ata executiva curta.

    Para CADA ata gerada, estruture OBRIGATORIAMENTE a resposta nos seguintes tópicos:

    ## 🏢 Ata de Reunião: [Nome da Empresa ou Cliente Identificado]

    Rapport:
    Contexto humano da reunião e nível de abertura do contato.

    Com quem eu estou falando:
    Cargo da pessoa, área de atuação e papel no processo decisório.

    O que o projeto está se encaminhando para ser:
    Síntese da oportunidade de projeto identificada. Problema a ser resolvido e solução proposta.

    Fatores cruciais de mapeamento:
    Elementos que impactam diretamente a viabilidade: sistemas existentes, estrutura de dados, cloud, processos, nível de maturidade digital.

    Stack Tecnológico (Glossário):
    - Liste aqui APENAS ferramentas, softwares, linguagens e infraestrutura citadas.
    - Use bullet points APENAS nesta seção.

    Termômetro do Lead (Scoring):
    Nota de 1 a 10 para o aquecimento do cliente e urgência.
    OBRIGATÓRIO colocar um destes emojis ao lado da nota para classificação visual: 
    🔴 (para notas de 1 a 4)
    🟡 (para notas de 5 a 7)
    🟢 (para notas de 8 a 10)
    Adicione 1 frase justificando a nota.

    Próximo passo claro:
    Próxima ação comercial objetiva.

    ---
    Regras importantes:
    - Se você identificar múltiplas reuniões diferentes, gere atas sequenciais.
    - Separe cada ata visualmente com uma linha (---).
    - Com EXCEÇÃO da seção "Stack Tecnológico", escreva em texto corrido, SEM listas ou bullet points.
    - A ata deve ser curta, clara e objetiva.

    Transcrição Bruta a ser analisada:
    {texto_transcricao}
    """

def criar_docx(texto_markdown):
    doc = Document()
    doc.add_heading('Diagnóstico Comercial - Relatório Gerado', 0)
    
    linhas = texto_markdown.split('\n')
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue
            
        if linha.startswith('## '):
            doc.add_heading(linha.replace('## ', ''), level=2)
        elif linha.startswith('**') and linha.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(linha.replace('**', '')).bold = True
        elif linha.startswith('- '):
            doc.add_paragraph(linha.replace('- ', ''), style='List Bullet')
        elif linha == '---':
            doc.add_page_break() 
        else:
            doc.add_paragraph(linha)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- 5. Interface de Usuário (UI) ---
opcoes_modelos = {
    "🟢 Básico: Gemini 2.5 Flash Lite (Rápido, custo menor)": "gemini-2.5-flash-lite",
    "🟡 Intermediário: Gemini 2.5 Flash (Padrão/Equilibrado)": "gemini-2.5-flash",
    "🔴 Pesado: Gemini 2.5 Pro (Maior raciocínio, mais lento)": "gemini-2.5-pro",
}

with st.sidebar:
    st.markdown("### ⚙️ Painel de Controlo")
    
    modelo_selecionado_nome = st.selectbox(
        "🧠 Escolha o Modelo de IA:",
        options=list(opcoes_modelos.keys()),
        index=1
    )
    modelo_id_real = opcoes_modelos[modelo_selecionado_nome]
    
    st.divider()
    st.button("🧹 Limpar Dados e Reiniciar", on_click=limpar_memoria, use_container_width=True)

st.title("💼 Assistente de Atas e Diagnóstico")
st.markdown("Extraia inteligência de reuniões com padronização executiva rigorosa e análise de viabilidade técnica.")
st.divider()

if client is None:
    st.error("🚨 Chave de API não encontrada nos Secrets.")
    st.stop()

# Layout de Upload
arquivos_pdf = st.file_uploader("📂 Faça o upload das transcrições (PDFs)", type=["pdf"], accept_multiple_files=True)

# Lógica de Empty State (Estado Vazio)
if not arquivos_pdf and not st.session_state.resultados_processados:
    st.markdown("""
        <div style="text-align: center; padding: 60px; background-color: #F8F9FA; border-radius: 15px; border: 2px dashed #CED4DA; margin-top: 20px;">
            <h2 style="color: #6C757D;">Nenhum documento carregado</h2>
            <p style="color: #6C757D; font-size: 16px;">Arraste os seus ficheiros PDF para a área acima ou clique para procurar.<br>A IA fará a leitura, separação de reuniões e análise de viabilidade técnica automaticamente.</p>
        </div>
    """, unsafe_allow_html=True)

# --- 6. Fluxo de Execução Principal ---
if arquivos_pdf and st.button("🚀 Processar Todos os Arquivos", type="primary"):
    with st.container():
        for arquivo in arquivos_pdf:
            with st.status(f"A processar {arquivo.name}...", expanded=True) as status:
                try:
                    st.write("Extraindo texto do PDF...")
                    texto_da_reuniao = extrair_texto_pdf(arquivo)
                    
                    if not texto_da_reuniao:
                        status.update(label=f"Falha: {arquivo.name} sem texto.", state="error")
                        continue
                    
                    st.write(f"Iniciando inferência com {modelo_id_real}...")
                    prompt_formatado = gerar_prompt_multiplas_atas(texto_da_reuniao)
                    
                    inicio_timer = time.time()
                    resposta = client.models.generate_content(
                        model=modelo_id_real,
                        contents=prompt_formatado
                    )
                    fim_timer = time.time()
                    
                    tokens_in = resposta.usage_metadata.prompt_token_count if hasattr(resposta, 'usage_metadata') and resposta.usage_metadata else 0
                    tokens_out = resposta.usage_metadata.candidates_token_count if hasattr(resposta, 'usage_metadata') and resposta.usage_metadata else 0
                    
                    st.session_state.resultados_processados[arquivo.name] = {
                        "texto_gerado": resposta.text,
                        "tokens_in": tokens_in,
                        "tokens_out": tokens_out,
                        "modelo_usado": modelo_id_real,
                        "tempo_segundos": (fim_timer - inicio_timer)
                    }
                    
                    status.update(label=f"Diagnóstico de {arquivo.name} concluído!", state="complete")
                    
                except Exception as e:
                    status.update(label="Erro no processamento", state="error")
                    st.error(f"Erro na API ({modelo_id_real}): {e}")
                    
    st.session_state.mostrar_animacao = True
    st.rerun()

# --- 7. Exibição, Edição e Governança ---
if st.session_state.resultados_processados:
    
    # Micro-interação de Sucesso
    if st.session_state.mostrar_animacao:
        st.balloons()
        st.toast("Todos os ficheiros foram processados com sucesso!", icon="✅")
        st.session_state.mostrar_animacao = False

    st.header("🎯 Resultados, Validação e Exportação")
    
    for nome_arquivo, dados in st.session_state.resultados_processados.items():
        with st.expander(f"⚙️ Gerenciar Ata: {nome_arquivo}", expanded=True):
            
            # Utilização de Cards (border=True) para organizar visualmente o texto
            st.markdown("##### 📝 Validação e Edição")
            st.info("Ata gerada com sucesso. O Termómetro do Lead foi codificado por cores automaticamente pela IA.")
            
            with st.container(border=True):
                texto_editado = st.text_area(
                    label="Área de edição:",
                    value=dados['texto_gerado'],
                    height=450,
                    key=f"edit_{nome_arquivo}",
                    label_visibility="collapsed"
                )
            
            st.markdown("##### 📥 Exportar Entregáveis")
            docx_file = criar_docx(texto_editado)
            
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.download_button("📄 Exportar para Word (.docx)", data=docx_file, file_name=f"Atas_{nome_arquivo.replace('.pdf', '')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx_{nome_arquivo}", use_container_width=True)
            with col_d2:
                st.download_button("📝 Exportar Texto Bruto (.txt)", data=texto_editado, file_name=f"Atas_{nome_arquivo.replace('.pdf', '')}.txt", mime="text/plain", key=f"txt_{nome_arquivo}", use_container_width=True)
            
            # --- Painel Nerd em formato de Cards ---
            mod = dados['modelo_usado']
            t_in = dados['tokens_in']
            t_out = dados['tokens_out']
            tempo = dados['tempo_segundos']
            
            custo_total = ((t_in / 1_000_000) * PRECOS_MODELOS[mod]["in"]) + ((t_out / 1_000_000) * PRECOS_MODELOS[mod]["out"])
            taxa_compressao = ((t_in - t_out) / t_in * 100) if t_in > 0 else 0
            tokens_por_segundo = (t_in + t_out) / tempo if tempo > 0 else 0

            st.markdown("<br>", unsafe_allow_html=True)
            with st.expander("🤓 Estatística para Nerd (Auditoria de Custo e IA)"):
                st.caption(f"**Modelo Utilizado:** `{mod}`")
                
                # Envolver as métricas num container com borda para dar aspeto de Dashboard
                with st.container(border=True):
                    col_t1, col_t2, col_t3 = st.columns(3)
                    col_t1.metric("Total de Tokens", f"{(t_in + t_out):,}")
                    col_t2.metric("Custo Estimado", f"${custo_total:.5f}")
                    col_t3.metric("Tempo de API", f"{tempo:.2f} s")
                    
                    st.divider()
                    
                    col_t4, col_t5, col_t6 = st.columns(3)
                    col_t4.metric("Compressão de Ruído", f"{taxa_compressao:.1f}%")
                    col_t5.metric("Velocidade", f"{tokens_por_segundo:,.0f} t/s")
                    col_t6.metric("Economia Estimada", f"~{(t_in / 300):.0f} min")

# --- Injeção de CSS Customizado para o File Uploader ---
st.markdown("""
    <style>
    /* Altera o fundo do card do arquivo carregado */
    [data-testid="stUploadedFile"] {
        background-color: #120421 !important; /* Usando o seu secondaryBackgroundColor */
        border: 1px solid #9D4EDD !important; /* Borda usando a sua primaryColor */
        border-radius: 8px;
    }
    
    /* Força os textos e ícones dentro do card a usarem a sua cor laranja */
    [data-testid="stUploadedFile"] div, 
    [data-testid="stUploadedFile"] span, 
    [data-testid="stUploadedFile"] svg {
        color: #f39334 !important; 
    }
    
    /* Remove o fundo branco do botão de fechar (X) do arquivo */
    button[title="Remove file"] {
        background-color: transparent !important;
    }
    button[title="Remove file"]:hover {
        background-color: #9D4EDD !important;
    }
    </style>
""", unsafe_allow_html=True)